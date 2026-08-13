#!/usr/bin/env python3
"""Insert Chapter 13 technical features into the editable software manual."""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


FONT = "Arial Unicode MS"

SECTIONS = (
    (
        "13.1  总体技术路线",
        "本软件采用 TypeScript、React 与 Vite 构建浏览器单页应用，将剧情播放、分支选择、回合制战斗、特殊玩法、存档和史乘展示组织在统一运行框架中。TypeScript 用于约束场景、战斗、剧情标记和存档数据结构，React 负责界面呈现，Vite 负责开发调试与生产构建。",
    ),
    (
        "13.2  部署方式",
        "浏览器版本构建后生成纯静态文件，可部署到静态 Web 服务器、对象存储或内容分发网络，终端通过现代浏览器直接访问，无须为单机剧情流程部署独立业务服务器。项目同时预留微信、抖音小游戏的平台适配层与独立构建入口；完成对应平台原生接口迁移后，可复用核心剧情、战斗和存档规则。",
    ),
    (
        "13.3  数据驱动与场景状态机",
        "剧情内容、场景跳转、分支结果和战斗参数采用结构化数据描述，主程序通过场景状态机决定当前界面及下一流程。新增剧情节点或战斗场景时，主要通过扩展数据配置与状态映射完成，减少在界面组件中重复编写流程判断。",
    ),
    (
        "13.4  业务逻辑与界面解耦",
        "战斗结算、教学战斗、撤离防守、存档兼容和剧情标记处理采用独立函数或业务模块实现，不直接依赖具体页面组件。与将数值计算、动画和界面事件集中在同一脚本的常见轻量游戏结构相比，本软件的界面层主要负责展示状态与转发操作，业务层负责计算结果，更便于测试、复用和后续平台迁移。",
    ),
    (
        "13.5  跨平台适配设计",
        "软件通过平台适配器统一封装存储、音频、图像、路由与分享能力。浏览器版本使用本地存储、HTML 音频、HTML 图像和历史路由；小游戏版本可替换为平台提供的对应接口。该方式避免在剧情和战斗模块中散布平台判断，使同一套核心规则能够服务不同发布形态。",
    ),
    (
        "13.6  本地存档与兼容控制",
        "游戏进度保存在当前设备浏览器本地，设置三个手动存档槽位、战前自动存档与战败回卷。存档包含版本号并区分兼容、旧版、未来版和无效数据，对未来版存档禁止覆盖，从而降低版本升级或异常数据导致进度损坏的风险。",
    ),
    (
        "13.7  资源加载与运行保障",
        "标题阶段预取主要背景、立绘和战斗资源，场景切换时复用缓存；单个图像或音频加载失败时采用降级处理，不阻断剧情推进。软件同时处理移动端首次交互后的音频解锁、响应式布局与减少动态效果偏好，以兼顾桌面端和移动端运行体验。",
    ),
    (
        "13.8  架构特点总结",
        "本软件的技术特点是以轻量静态部署承载单机历史叙事体验，以数据驱动状态机组织长流程，以独立业务内核保证战斗与存档规则一致，并通过平台适配层预留多端发布能力。相较依赖实时服务端、账号体系或将全部逻辑耦合在页面脚本中的游戏架构，该方案部署成本较低、核心规则边界清晰，适合章节内容持续扩展。",
    ),
)


def set_font(run, size: float, *, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def set_text(paragraph, text: str, *, size: float, bold: bool) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run, size, bold=bold)


def clone_after(paragraph):
    clone = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(clone)
    return clone


def update_toc(document: Document) -> None:
    chapter12 = next(p for p in document.paragraphs if p.text.strip().startswith("第12章") and "\t12" in p.text)
    clone = clone_after(chapter12)
    texts = clone.xpath(".//w:t")
    if not texts:
        raise RuntimeError("Unable to clone table-of-contents entry")
    texts[0].text = "第13章  技术特点\t24"
    for extra in texts[1:]:
        extra.text = ""

    replacements = {
        "附录A—M": ("附录A—M  设计数据与实现明细", 25),
        "附录N": ("附录N  源码模块与依赖追溯", 34),
        "附录O": ("附录O  类型、函数与常量索引", 43),
    }
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        for prefix, (label, page) in replacements.items():
            if value.startswith(prefix) and "\t" in paragraph.text:
                set_text(paragraph, f"{label}\t{page}", size=10.5, bold=False)
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("说明：目录页码包含封面和目录"):
            set_text(
                paragraph,
                "说明：目录页码包含封面和目录；正文第3页起，第12章运行截图第12页起，第13章技术特点第24页，附录第25页起。",
                size=8.5,
                bold=False,
            )


def insert_chapter(document: Document) -> None:
    if any(p.text.strip().startswith("第13章") and "\t" not in p.text for p in document.paragraphs):
        raise RuntimeError("Chapter 13 already exists")
    appendix = next(p for p in document.paragraphs if p.text.strip().startswith("附录A  场景状态全表"))

    chapter = appendix.insert_paragraph_before()
    chapter.paragraph_format.space_before = Pt(0)
    chapter.paragraph_format.space_after = Pt(7)
    chapter.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    chapter.paragraph_format.line_spacing = Pt(21)
    set_text(chapter, "第13章  技术特点", size=13, bold=True)

    for heading, body in SECTIONS:
        p = appendix.insert_paragraph_before()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)
        set_text(p, heading, size=10.5, bold=True)

        p = appendix.insert_paragraph_before()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(17.5)
        p.paragraph_format.first_line_indent = Mm(7.4)
        set_text(p, body, size=9.6, bold=False)

    page_break = appendix.insert_paragraph_before()
    page_break.paragraph_format.space_before = Pt(0)
    page_break.paragraph_format.space_after = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    document = Document(args.docx)
    update_toc(document)
    insert_chapter(document)
    document.save(args.docx)


if __name__ == "__main__":
    main()
