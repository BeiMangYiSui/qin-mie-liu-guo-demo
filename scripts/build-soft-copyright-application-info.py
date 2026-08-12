#!/usr/bin/env python3
"""Build the applicant-facing software copyright filing aids.

Sensitive applicant fields are read from existing local release documents at
runtime. They are deliberately not embedded in this script or printed to
stdout.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


REPO = Path("/Users/zhuangxiji/Desktop/CC/qin-mie-liu-guo-demo")
PACKAGE = Path("/Users/zhuangxiji/Desktop/秦灭六国软著申请/Codex-最终材料-20260812")
PERSONAL = REPO / "docs/release/抖音上线待做事项/个人基本信息.md"
FILING = REPO / "docs/release/微信上线待做事项/02-小程序备案信息预填表.md"
OUT = PACKAGE / "03-申请信息与提交指引-含敏感信息"
MISSING = PACKAGE / "05-本人必须补充-勿上传空白说明"

SOFTWARE = "秦灭六国游戏软件"
SHORT_NAME = "秦灭六国"
VERSION = "V1.0"
COMPLETION_DATE = "2026年08月05日"
FONT = "Arial Unicode MS"


@dataclass(frozen=True)
class Applicant:
    name: str
    id_type: str
    id_number: str
    address: str
    phone: str
    email: str


def read_colon(text: str, label: str) -> str:
    match = re.search(rf"^\s*{re.escape(label)}\s*[：:]\s*(.+?)\s*$", text, re.M)
    if not match:
        raise ValueError(f"Missing field: {label}")
    return match.group(1).strip().strip("`*")


def read_table(text: str, label: str) -> str:
    for line in text.splitlines():
        if "|" not in line:
            continue
        cells = [cell.strip().strip("`*") for cell in line.strip().strip("|").split("|")]
        for index, cell in enumerate(cells[:-1]):
            if cell == label:
                return cells[index + 1]
    raise ValueError(f"Missing table field: {label}")


def validate_prc_id(value: str) -> None:
    if not re.fullmatch(r"\d{17}[\dXx]", value):
        raise ValueError("Applicant ID number is not an 18-character PRC ID")
    weights = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
    checks = "10X98765432"
    expected = checks[sum(int(n) * w for n, w in zip(value[:17], weights)) % 11]
    if value[-1].upper() != expected:
        raise ValueError("Applicant ID checksum is invalid")


def load_applicant() -> Applicant:
    personal = PERSONAL.read_text(encoding="utf-8")
    filing = FILING.read_text(encoding="utf-8")
    name = read_colon(personal, "姓名")
    email = read_colon(personal, "邮箱")
    id_number = read_colon(personal, "身份证号")
    filing_name = read_table(filing, "主办者姓名")
    if filing_name != name:
        raise ValueError("Applicant names disagree between local source files")
    id_type_raw = read_table(filing, "证件类型")
    filing_id_raw = read_table(filing, "证件号码")
    filing_id = re.sub(r"[^0-9Xx]", "", filing_id_raw)
    # The filing template may intentionally say "本人按身份证填写". Compare only
    # when that file contains an actual ID number.
    if filing_id and filing_id != id_number:
        raise ValueError("Applicant ID numbers disagree between local source files")
    phone = re.sub(r"\s+", "", read_table(filing, "手机号码"))
    filing_email = read_table(filing, "电子邮箱")
    if filing_email != email:
        raise ValueError("Applicant emails disagree between local source files")
    address = read_table(filing, "住所/通信地址")
    validate_prc_id(id_number)
    if not re.fullmatch(r"1\d{10}", phone):
        raise ValueError("Applicant phone number is not an 11-digit mobile number")
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
        raise ValueError("Applicant email format is invalid")
    id_type = "居民身份证" if "身份证" in id_type_raw else id_type_raw
    return Applicant(name, id_type, id_number, address, phone, email)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=9.5, bold=False, color="1F2937") -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— 第 ")
    set_run(run, 8, color="6B7280")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页 —")
    set_run(run, 8, color="6B7280")


def configure(doc: Document, title: str) -> None:
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(16)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(17)
    sec.right_margin = Mm(17)
    sec.header_distance = Mm(8)
    sec.footer_distance = Mm(8)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.1
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{SOFTWARE} {VERSION}｜{title}")
    set_run(run, 8, color="6B7280")
    add_page_number(sec.footer.paragraphs[0])
    props = doc.core_properties
    props.title = title
    props.subject = f"{SOFTWARE} {VERSION} 软件著作权登记填报辅助材料"
    props.author = "申请人"


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, 18, bold=True, color="17365D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle)
    set_run(r, 10, color="4B5563")


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, 12 if level == 1 else 10.5, bold=True, color="1F4E78")


def add_notice(doc: Document, text: str, color="FFF2CC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run(r, 9, bold=True, color="7F6000")


def add_fields_table(doc: Document, rows: list[tuple[str, str, str]], widths=(35, 90, 45)) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for index, text in enumerate(("登记字段", "建议填写内容", "填写说明/依据")):
        cell = table.rows[0].cells[index]
        cell.width = Mm(widths[index])
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run(r, 9, bold=True, color="17365D")
    for label, value, note in rows:
        cells = table.add_row().cells
        for i, text in enumerate((label, value, note)):
            cells[i].width = Mm(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, 8.8, bold=(i == 0), color="1F2937")
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_bullets(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Mm(6)
        p.paragraph_format.first_line_indent = Mm(-3)
        r = p.add_run(text)
        set_run(r, 9.2)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break()
    paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run = OxmlElement("w:r")
    run.append(page_break)
    paragraph._p.append(run)


def build_application_info(applicant: Applicant, path: Path) -> None:
    doc = Document()
    configure(doc, "登记申请信息最终填写稿")
    add_title(doc, "计算机软件著作权登记申请信息最终填写稿", f"{SOFTWARE}｜{VERSION}｜自然人申请")
    add_notice(
        doc,
        "本文件含身份证号、电话、住址等敏感信息，仅供本人照录登记系统。它不是版权保护中心生成的官方申请表，不作为鉴别材料上传。",
    )
    add_heading(doc, "一、软件基本信息")
    add_fields_table(doc, [
        ("软件全称", SOFTWARE, "与两份鉴别材料标题一致"),
        ("软件简称", SHORT_NAME, "项目及游戏标题使用名称"),
        ("版本号", VERSION, "使用大写 V；不要改成 v1.0"),
        ("软件分类", "游戏软件", "若系统为下拉选项，选择最接近的游戏软件类别"),
        ("开发方式", "独立开发", "按现有个人开发及发布资料预填"),
        ("开发完成日期", COMPLETION_DATE, "以 V1.0 代码与发布材料冻结日为依据"),
        ("发表状态", "未发表", "公开 Pages 地址当前为 404，发布材料记载尚未上线"),
        ("首次发表日期/地点", "不适用", "未发表时不填写发表日期和地点"),
        ("权利取得方式", "原始取得", "独立开发形成"),
        ("权利范围", "全部权利", "无转让、继承或部分权利材料"),
        ("是否修改他人软件", "否", "项目为独立开发，不按改编软件申报"),
    ])

    add_page_break(doc)
    add_heading(doc, "二、著作权人及联系人信息")
    add_fields_table(doc, [
        ("著作权人类型", "自然人", "个人申请"),
        ("姓名", applicant.name, "须与身份证完全一致"),
        ("国籍/地区", "中国", "居民身份证持有人"),
        ("证件类型", applicant.id_type, "身份证正反面 JPG 已准备"),
        ("证件号码", applicant.id_number, "已通过 18 位身份证校验规则"),
        ("通信地址", applicant.address, "来自现有备案预填资料；提交前本人核对门牌信息"),
        ("邮政编码", "401122", "两江新区政府公布的悦来街道嘉悦社区邮编"),
        ("联系电话", applicant.phone, "来自现有备案预填资料"),
        ("电子邮箱", applicant.email, "来自现有个人基本信息"),
        ("联系人", applicant.name, "本人办理"),
        ("代理人", "无", "本人在线办理时选择无代理"),
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(doc, "软件技术信息填写稿", "用于登记系统中的开发、功能和技术特点字段")
    add_heading(doc, "三、开发与运行环境")
    add_fields_table(doc, [
        ("开发硬件环境", "Apple 芯片 Mac 计算机、通用显示器及输入设备", "不填写未实测的内存或性能数据"),
        ("开发操作系统", "macOS 15", "当前项目开发环境"),
        ("开发工具", "Node.js 24、npm 11、TypeScript、Vite、ESLint", "与项目工具链一致"),
        ("运行硬件环境", "通用桌面计算机或移动智能终端", "浏览器运行形态"),
        ("运行操作系统", "Windows、macOS、iOS 或 Android", "需具备现代浏览器"),
        ("运行支撑软件", "Chrome、Edge、Safari 等现代浏览器", "当前 V1.0 不申报为原生小游戏"),
        ("编程语言", "TypeScript、JavaScript、TSX、CSS", "核心业务以 TypeScript/TSX 编写"),
        ("源程序量", "17057 行", "src 目录 128 个源文件的物理行数"),
    ])
    add_heading(doc, "四、开发目的")
    p = doc.add_paragraph()
    r = p.add_run(
        "通过章节化历史叙事、剧情分支选择和轻量回合制战斗，为用户提供单人策略剧情体验；将剧情阅读、行动决策、战斗目标、任务结果与历史对照内容组织在统一流程中，并通过本地存档保存用户进度和关键选择。"
    )
    set_run(r, 9.5)
    add_heading(doc, "五、面向领域/行业")
    p = doc.add_paragraph()
    r = p.add_run("数字文化娱乐、游戏软件。")
    set_run(r, 9.5)
    add_heading(doc, "六、主要功能")
    p = doc.add_paragraph()
    r = p.add_run(
        "本软件是一款以战国末期为背景的单人历史叙事策略游戏，提供章节剧情播放、人物对话、关键分支选择、轻量回合制战斗、教学战斗、火场救援、残军追截、疫营撤离、章末结算、史乘对照和本地存档等功能。用户选择和战斗结果通过剧情状态标记保存，并影响后续场景、战斗配置和结算内容。软件还提供背景音乐、角色配音、操作音效、媒体资源预加载、移动端音频解锁及存档版本兼容处理。"
    )
    set_run(r, 9.5)
    add_heading(doc, "七、技术特点")
    p = doc.add_paragraph()
    r = p.add_run(
        "软件采用 TypeScript、React 和 Vite 开发，使用组件化界面结构，并将剧情数据、战斗规则、存档逻辑和界面渲染分层组织。核心战斗与存档逻辑采用类型化数据结构管理，支持多种战斗目标、剧情状态标记、存档兼容性检测和异常数据过滤。当前版本运行于现代浏览器环境，游戏进度保存在用户设备本地，不依赖业务服务器保存个人游戏数据。"
    )
    set_run(r, 9.5)

    add_heading(doc, "八、权属声明填写口径")
    add_bullets(doc, [
        "开发性质：个人独立开发；无合作开发者、委托方或职务开发单位。",
        "权利来源：原始取得；申请全部权利。",
        "软件不是在他人既有软件基础上修改形成，不提交原软件著作权人许可证明。",
        "如实际存在未记载的共同开发、委托、任职单位任务或权利转让情况，必须停止使用上述口径并补充协议或证明。",
    ])
    add_heading(doc, "九、提交前本人逐项确认")
    add_bullets(doc, [
        f"软件全称和版本号确认为“{SOFTWARE}”“{VERSION}”。",
        f"开发完成日期确认为 {COMPLETION_DATE}；如真实完成日不同，先修改登记系统再提交。",
        "截至申请日软件未向不特定公众开放；如曾有可公开访问的正式版本，须改填已发表并写实际日期、地点。",
        "著作权人、证件号码、电话和邮箱均与本人当前有效资料一致；通信地址需补充门牌及邮编。",
    ])
    add_notice(doc, "提交前最后核对：姓名、证件号码、地址、电话、完成日期及未发表状态均属于法律事实；若现实情况与本稿不同，以真实情况为准并同步修改登记系统。", "FCE4D6")
    doc.save(path)


def build_submission_guide(applicant: Applicant, path: Path) -> None:
    doc = Document()
    configure(doc, "提交顺序与文件清单")
    add_title(doc, "软件著作权登记提交顺序与文件清单", f"{SOFTWARE}｜{VERSION}")
    add_notice(doc, "两份技术 PDF 和身份证正反面 JPG 已准备并通过校验。本人仍需补全通信地址门牌与邮编，并在登记系统生成、签署和上传 R11 申请确认签章页。")
    add_heading(doc, "一、登记系统照录信息")
    add_fields_table(doc, [
        ("软件", f"{SOFTWARE} {VERSION}", "不得改名或改版本号"),
        ("申请人", applicant.name, "自然人、本人办理"),
        ("开发方式", "独立开发", "按现有项目资料预填"),
        ("完成日期", COMPLETION_DATE, "V1.0 冻结日"),
        ("发表状态", "未发表", "不填写首次发表日期/地点"),
        ("取得方式/范围", "原始取得 / 全部权利", "无权利转让或继承"),
    ])
    add_heading(doc, "二、上传文件")
    add_fields_table(doc, [
        ("程序鉴别材料", "01-秦灭六国游戏软件-V1.0-源程序鉴别材料.pdf", "60 页，A4，每页 50 行"),
        ("文档鉴别材料", "02-秦灭六国游戏软件-V1.0-软件设计与使用说明书.pdf", "45 页，A4，每页不少于 30 行"),
        ("身份证明", "01-居民身份证正面.jpg / 02-居民身份证反面.jpg", "已准备；官方页面按正反面分栏上传"),
        ("权属证明", "通常无需另交", "独立开发、原始取得且无合作/委托/职务情形时"),
        ("申请确认签章页", "登记系统生成后下载", "自然人本人签名并填写身份证号码，再上传签章原件"),
    ])
    add_heading(doc, "三、操作顺序")
    add_bullets(doc, [
        "登录中国版权保护中心的软件登记业务入口，创建“计算机软件著作权登记”申请。",
        "按《登记申请信息最终填写稿》逐项照录；遇到系统字段名称变化，按语义选择对应项。",
        "上传源程序 PDF、软件设计与使用说明书 PDF，以及已准备的身份证正面和反面 JPG。",
        "系统生成申请确认签章页后，逐项核对软件名称、版本、申请人、完成日期和发表状态；本人签名并填写身份证号码后上传签章原件。",
        "按系统当时提示完成电子确认、签名/盖章及后续提交；保留受理号和最终上传文件副本。",
    ])
    add_heading(doc, "四、禁止上传")
    add_bullets(doc, [
        "本《提交顺序与文件清单》和《登记申请信息最终填写稿》仅供本人操作，不作为鉴别材料上传。",
        "QA 联系表、追溯 CSV、构建元数据、可编辑 DOCX、旧版 Qoder 报告和重复截图均不上传。",
        "不得上传包含访问令牌、密钥、身份证号明文的项目配置或仓库截图。",
    ])
    add_heading(doc, "五、提交后留存")
    add_bullets(doc, [
        "保存登记系统生成的最终申请表、受理号、提交时间和全部上传文件；不要仅保留浏览器截图。",
        "保存本包 `04-QA与校验/SHA256SUMS.txt`，用于证明提交技术 PDF 与本地归档文件一致。",
        "收到补正通知时，只针对通知列出的项目修改；软件名称、版本、完成日期或著作权人发生变化时，同时复核所有材料的一致性。",
    ])
    add_heading(doc, "六、当前材料状态")
    add_fields_table(doc, [
        ("源程序 PDF", "已完成并通过 QA", "60 页、每页 50 行"),
        ("说明书 PDF", "已完成并通过 QA", "45 页、每页不少于 30 行"),
        ("申请信息填写稿", "已预填", "含敏感信息，仅本人使用"),
        ("身份证明", "已完成并核对", "正面姓名/号码匹配，反面类型匹配"),
        ("通信地址/邮编", "待本人补充", "现有资料只到小区级"),
        ("申请确认签章页", "待系统生成并本人签名", "需登录本人登记账号"),
    ])
    doc.save(path)


def write_markdown_files(applicant: Applicant) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    MISSING.mkdir(parents=True, exist_ok=True)
    (MISSING / "README-本人必须补充.md").write_text(
        f"""# 本人必须补充的两项信息/操作

身份证正反面已在私密目录中准备完成。以下内容仍无法由项目源码或现有文档代替：

1. **通信地址门牌信息**。现有资料只到小区级；请在本目录 `通信地址门牌信息-本人本地填写.txt` 中补充楼栋、单元、楼层和房号。官方页面的收件地址规范明确要求填写到房间号。
2. **R11 申请确认签章页**。登录系统照录 `../03-申请信息与提交指引-含敏感信息/01-登记申请信息最终填写稿.pdf`。系统生成签章页后，自然人申请人本人签名并填写身份证号码，再按页面要求上传签章原件。

本目录中的说明文件不能代替本人签名和登记系统生成的签章页。技术鉴别材料与身份证明已经准备完成。
""",
        encoding="utf-8",
    )
    (OUT / "README-敏感信息说明.md").write_text(
        """# 敏感信息说明

本目录文件含申请人姓名、身份证号码、通信地址、电话和电子邮箱，仅供本人办理软件著作权登记使用。请勿发送到公开群聊、代码仓库、网盘公开链接或与申请无关的第三方。

`01-登记申请信息最终填写稿` 是在线填写辅助稿，不是官方申请表；`02-提交顺序与文件清单` 是操作清单，二者均不作为软件鉴别材料上传。
""",
        encoding="utf-8",
    )


def main() -> None:
    applicant = load_applicant()
    OUT.mkdir(parents=True, exist_ok=True)
    info = OUT / "01-登记申请信息最终填写稿.docx"
    guide = OUT / "02-提交顺序与文件清单.docx"
    build_application_info(applicant, info)
    build_submission_guide(applicant, guide)
    write_markdown_files(applicant)
    print("Built applicant filing aids; sensitive values were not printed.")


if __name__ == "__main__":
    main()
