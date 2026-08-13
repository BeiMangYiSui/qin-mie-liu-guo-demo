#!/usr/bin/env python3
"""Normalize the Chapter 13 table-of-contents entry after insertion."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


FONT = "Arial Unicode MS"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    document = Document(args.docx)
    paragraph = next(
        p
        for p in document.paragraphs[:40]
        if p.text.strip().startswith("第13章")
    )
    paragraph.clear()
    run = paragraph.add_run("第13章  技术特点\t24")
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    document.save(args.docx)


if __name__ == "__main__":
    main()
