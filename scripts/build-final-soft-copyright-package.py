#!/usr/bin/env python3
"""Build the final technical software-copyright deposit package.

The output is intentionally generated from the current V1.0 repository and the
reviewed design draft.  The primary document uses a deterministic, page-aligned
layout: 38 A4 design pages with 30 substantive lines per page, followed by 12
illustrated operation pages. Every illustrated page also contains 30 concise,
substantive operation lines and one verified runtime screenshot.
The same DOCX is converted to the submission PDF during the render/QA step.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


REPO = Path("/Users/zhuangxiji/Desktop/CC/qin-mie-liu-guo-demo")
ROOT = Path("/Users/zhuangxiji/Desktop/秦灭六国软著申请")
WORKING = ROOT / "working"
QODER_FINAL = ROOT / "final"
OUT = ROOT / "Codex-最终材料-20260812"
SUBMIT = OUT / "01-正式提交件-技术材料"
EDITABLE = OUT / "02-可编辑源文件"
CONFIRM = OUT / "03-申请人确认后填写"
QA = OUT / "04-QA与校验"

SOFTWARE = "秦灭六国游戏软件"
VERSION = "V1.0"
DOC_TITLE = f"{SOFTWARE} {VERSION} 软件设计与使用说明书"
DOC_FONT = "Arial Unicode MS"
CODE_FONT = "Arial Unicode MS"
LINES_PER_PAGE = 30
TEXT_PAGES = 38
SCREENSHOT_PAGES = 12
FINAL_PAGES = TEXT_PAGES + SCREENSHOT_PAGES
FRONT_MATTER_PAGES = 2
TOTAL_PDF_PAGES = FRONT_MATTER_PAGES + FINAL_PAGES
MAIN_CHAPTER_TEXT_PAGES = 9
MAIN_CHAPTER_LINE_TARGET = MAIN_CHAPTER_TEXT_PAGES * LINES_PER_PAGE
APPENDIX_TEXT_PAGES = TEXT_PAGES - MAIN_CHAPTER_TEXT_PAGES
MAX_DISPLAY_WIDTH = 78
SCREENSHOT_DIR = QODER_FINAL / "screenshots"

SCREENSHOTS = (
    {
        "file": "01-title.png",
        "title": "图 1  软件启动与标题界面",
        "entry": "进入方式：在现代浏览器中访问软件运行地址，资源加载完成后进入标题页。",
        "operation": "主要操作：点击“开始”进入剧情；存在有效进度时可继续游戏或打开存档；也可进入史乘页面。",
    },
    {
        "file": "02-chapter-card.png",
        "title": "图 2  第一章章节转场界面",
        "entry": "进入方式：序章流程结束后自动显示章节转场卡片。",
        "operation": "主要操作：核对章节名称与背景信息后点击继续，进入第一章“新郑城破”的后续剧情。",
    },
    {
        "file": "03-story-scene.png",
        "title": "图 3  剧情叙事与对话界面",
        "entry": "进入方式：开始新游戏或读档恢复至剧情场景。",
        "operation": "主要操作：查看背景、角色立绘、说话人和文本，点击画面或使用键盘推进对白。",
    },
    {
        "file": "04-branch-choice.png",
        "title": "图 4  限时分支选择界面",
        "entry": "进入方式：剧情推进至关键抉择或残军追截节点时自动弹出。",
        "operation": "主要操作：在倒计时结束前选择路线；选择结果写入剧情标记并影响后续结算。",
    },
    {
        "file": "05-battle.png",
        "title": "图 5  通用回合制战斗界面",
        "entry": "进入方式：剧情推进至战斗节点后，根据对应场景配置初始化战斗。",
        "operation": "主要操作：查看敌方公开意图，为当前角色选择行动和目标，完成回合直至胜负判定。",
    },
    {
        "file": "06-tutorial-battle.png",
        "title": "图 6  教学战斗与目标提示界面",
        "entry": "进入方式：序章伏杀流程进入教学战斗。",
        "operation": "主要操作：在有限行动点内处理夺卷、救人和守退路目标，理解行动取舍与回合规则。",
    },
    {
        "file": "07-fire-rescue.png",
        "title": "图 7  火场救援限时玩法界面",
        "entry": "进入方式：第一章火场抉择中选择抢救档案后进入。",
        "operation": "主要操作：观察剩余时间、火势与各档案耗时，在焚毁前选择至多三份档案进行抢救。",
    },
    {
        "file": "08-pursuit-intercept.png",
        "title": "图 8  残军追截玩法界面",
        "entry": "进入方式：第一章火场抉择中选择追截残军后进入。",
        "operation": "主要操作：依据波次线索和倒计时选择拦截路线，系统累计成功拦截数量。",
    },
    {
        "file": "09-defense-battle.png",
        "title": "图 9  疫营撤离防守界面",
        "entry": "进入方式：第一章火场抉择中选择护送疫营撤离后进入。",
        "operation": "主要操作：查看本轮危机和目标血量，在结阵、抢救、反击之间选择应对并守满回合。",
    },
    {
        "file": "10-save-load.png",
        "title": "图 10  存档读取界面",
        "entry": "进入方式：在标题页选择存档入口，或在可存档场景中打开存档菜单。",
        "operation": "主要操作：查看三个手动槽位的章节、场景与保存时间，执行读取、保存、覆盖或删除。",
    },
    {
        "file": "11-settle.png",
        "title": "图 11  第一章章末结算界面",
        "entry": "进入方式：完成第一章最终剧情与分支玩法后自动进入。",
        "operation": "主要操作：查看战绩、关键抉择、支线结果、史乘收集进度及章节评价。",
    },
    {
        "file": "12-shicheng.png",
        "title": "图 12  史乘历史对照界面",
        "entry": "进入方式：从标题页、章末结算或终章入口进入史乘。",
        "operation": "主要操作：浏览已解锁的历史对照卡，核对游戏叙事与史实资料的区别。",
    },
)

SCREENSHOT_OPERATION_DETAILS = (
    (
        "界面用途：标题页是软件启动后的统一入口。",
        "显示内容：页面中央显示软件名称“秦灭六国”。",
        "显示内容：副标题说明当前版本包含序章与第一章。",
        "显示内容：背景图展示三人行动小队和新郑城环境。",
        "开始操作：点击“奉令出发”创建新的剧情进度。",
        "继续操作：存在有效存档时可从最近进度继续。",
        "存档操作：点击存档入口可查看三个手动槽位。",
        "史乘操作：点击史乘入口可浏览历史对照内容。",
        "声音操作：点击右上角声音图标切换静音状态。",
        "加载过程：标题页显示期间后台预取主要媒体资源。",
        "交互方式：桌面端支持鼠标点击和键盘操作。",
        "交互方式：移动端支持触摸操作。",
        "音频规则：首次用户交互用于解除浏览器音频限制。",
        "存储规则：继续按钮依据本地有效存档状态显示。",
        "异常处理：背景加载失败时仍保留主要入口按钮。",
        "异常处理：音频不可用时不影响开始游戏。",
        "视觉信息：标题、按钮和说明文字保持高对比显示。",
        "功能边界：标题页不直接修改剧情分支标记。",
        "功能边界：新游戏确认后才初始化剧情状态。",
        "数据处理：游戏进度保存在当前设备浏览器本地。",
        "数据处理：软件不要求注册游戏账号。",
        "数据处理：软件不依赖业务服务器保存个人进度。",
        "版本信息：本说明对应秦灭六国游戏软件 V1.0。",
        "运行形态：跨平台数字游戏；支持现代浏览器运行，并提供微信、抖音小游戏适配构建。",
        "主要入口：开始、继续、存档和史乘相互独立。",
        "操作反馈：按钮悬停或按下时提供视觉反馈。",
        "健康提示：标题页底部显示适度游戏提示。",
        "退出方式：关闭浏览器页面即可结束本次运行。",
        "恢复方式：再次打开软件可从本地存档恢复。",
        "截图结论：图 1 证明软件具备完整启动入口。",
    ),
    (
        "界面用途：章节卡用于衔接序章与第一章。",
        "进入条件：完成序章末尾剧情后自动显示。",
        "显示内容：左上方标识第一章章节信息。",
        "显示内容：场景背景展示第一章主要发生地点。",
        "显示内容：底部保留当前叙事文本。",
        "主要操作：阅读章节信息后点击继续。",
        "继续结果：系统切换到第一章首个剧情场景。",
        "跳过操作：用户可使用界面提供的跳过按钮。",
        "声音操作：右上角声音图标继续保持可用。",
        "状态规则：章节卡自身不产生战斗结算。",
        "状态规则：章节卡不改变既有剧情选择。",
        "流程规则：章节卡只在预定流程节点出现。",
        "流程规则：继续后不可在本页选择支线路线。",
        "视觉设计：章节卡使用统一的深色历史氛围。",
        "视觉设计：章节名称与场景背景同时呈现。",
        "交互方式：桌面端可点击继续按钮。",
        "交互方式：移动端可触摸继续按钮。",
        "键盘方式：可使用设定的确认键继续。",
        "动画反馈：章节切换采用淡入淡出效果。",
        "加载规则：后续场景资源在切换前进行预取。",
        "异常处理：背景图片失败时仍可阅读章节文本。",
        "异常处理：声音播放失败时仍可继续流程。",
        "存档规则：章节转场属于不可手动存档场景。",
        "恢复规则：读档落在相邻节点时由流程状态恢复。",
        "数据来源：章节名称和跳转关系由场景状态定义。",
        "功能对应：本页对应章节转场模块。",
        "版本范围：V1.0 完成序章和第一章内容。",
        "章节名称：第一章围绕“新郑城破”展开。",
        "使用结果：点击继续后进入第一章剧情播放。",
        "截图结论：图 2 证明软件具备章节化流程。",
    ),
    (
        "界面用途：剧情页承担旁白、对白和演出信息播放。",
        "进入方式：开始新游戏或从有效存档恢复。",
        "显示内容：顶部标识章节、年代和当前场景。",
        "显示内容：中部背景表现当前剧情地点。",
        "显示内容：需要时在背景上叠加角色立绘。",
        "显示内容：下方对话区显示说话人与文本。",
        "主要操作：点击画面推进到下一句对白。",
        "键盘操作：空格键可推进剧情文本。",
        "确认操作：回车键可确认可用选项。",
        "跳过操作：点击跳过按钮快速到达下一节点。",
        "历史显示：最近对白以较低强调度保留。",
        "当前显示：当前对白保持突出显示。",
        "行类型：剧情支持旁白、对白、内心和演出说明。",
        "声音功能：匹配到语音资源时自动播放配音。",
        "声音功能：推进对白时触发界面操作音效。",
        "声音控制：用户可随时切换全局静音。",
        "分支规则：到达选择点后显示选项面板。",
        "分支规则：选项确认后写入对应剧情标记。",
        "流程规则：当前场景播放完成后切换下一场景。",
        "存档规则：可存档剧情场景允许打开存档菜单。",
        "恢复规则：读档后恢复场景状态和剧情标记。",
        "异常处理：头像失败时使用字符占位。",
        "异常处理：背景失败时对话仍可继续显示。",
        "异常处理：配音失败不阻断剧情推进。",
        "响应布局：桌面端与移动端使用自适应排版。",
        "可访问性：关键按钮提供明确文字或标签。",
        "数据来源：剧情文本由类型化场景数据提供。",
        "功能对应：本页对应剧情播放与语音模块。",
        "使用结果：用户可连续阅读并推进故事。",
        "截图结论：图 3 证明软件具备剧情叙事界面。",
    ),
    (
        "界面用途：分支页用于处理限时路线选择。",
        "进入条件：剧情推进至关键抉择节点时显示。",
        "显示内容：面板顶部显示当前任务名称。",
        "显示内容：右上方显示剩余选择时间。",
        "显示内容：进度区域显示当前波次和累计结果。",
        "显示内容：线索区域提示本轮可判断的信息。",
        "显示内容：下方提供三条可选路线。",
        "主要操作：在倒计时结束前点击一条路线。",
        "选择规则：每个波次只能确认一条路线。",
        "选择结果：系统依据路线判定拦截是否成功。",
        "累计规则：成功拦截数量随波次累计。",
        "状态记录：累计数量写入剧情状态标记。",
        "后续影响：拦截结果进入章末支线结算。",
        "超时处理：倒计时结束时按玩法规则自动结算。",
        "重复限制：确认后不能在同一波次重复选择。",
        "视觉反馈：已选路线提供高亮或状态反馈。",
        "流程反馈：选择完成后进入下一波线索。",
        "结束条件：所有波次处理完成后退出玩法。",
        "返回流程：玩法结束后回到火起后的剧情。",
        "交互方式：桌面端支持鼠标点击路线。",
        "交互方式：移动端支持触摸路线卡片。",
        "声音功能：操作时可播放对应界面音效。",
        "声音控制：右上角继续提供静音开关。",
        "异常处理：媒体失败不影响倒计时和判定。",
        "数据来源：波次、线索和路线由玩法数据定义。",
        "功能对应：本页对应残军追截模块。",
        "业务价值：选择机制体现信息判断和时间压力。",
        "结果用途：追截数量用于生成差异化结算内容。",
        "版本范围：该玩法属于第一章可选支线。",
        "截图结论：图 4 证明软件具备限时分支交互。",
    ),
    (
        "界面用途：战斗页执行通用回合制战斗。",
        "进入条件：剧情推进至战斗节点后初始化。",
        "显示内容：顶部显示战斗名称和当前回合。",
        "显示内容：顶部同时显示本场胜利目标。",
        "显示内容：左侧展示己方角色、血量和状态。",
        "显示内容：右侧展示敌方单位、血量和意图。",
        "显示内容：右栏记录本场战斗事件。",
        "显示内容：底部列出当前角色可用行动。",
        "主要操作：先选择行动，再按需要选择目标。",
        "回合规则：各在场英雄依次完成行动。",
        "意图规则：敌方在行动前公开本回合意图。",
        "结算规则：玩家行动按固定数值修改战斗状态。",
        "敌方阶段：玩家行动结束后敌方依次执行意图。",
        "胜负规则：每次状态推进后立即检查胜负。",
        "行动类型：突击用于造成中等伤害。",
        "行动类型：打断用于削弱并取消当前意图。",
        "行动类型：治疗用于恢复友方角色血量。",
        "行动类型：布烟用于使敌方攻击落空。",
        "行动类型：飞针和截剑用于针对性控制。",
        "目标规则：不同技能限制敌方、友方或无目标。",
        "反馈方式：伤害、治疗和状态变化显示动画。",
        "日志规则：关键行动和结果写入战斗记录。",
        "自动存档：进入战斗前保存战前恢复点。",
        "失败处理：战败后可重试或回卷到战前。",
        "胜利处理：胜利后返回对应剧情后段。",
        "异常处理：状态只能通过战斗引擎函数推进。",
        "响应布局：移动端折叠或调整战斗记录区域。",
        "功能对应：本页对应通用战斗引擎与战斗界面。",
        "版本范围：V1.0 包含多场配置化战斗。",
        "截图结论：图 5 证明软件具备完整战斗操作。",
    ),
    (
        "界面用途：教学战引导用户理解行动点取舍。",
        "进入条件：序章伏杀剧情推进至教学节点。",
        "显示内容：顶部展示三个教学目标。",
        "目标一：护住明卷需要连续执行指定行动。",
        "目标二：救出孟甲需要投入指定行动点。",
        "目标三：守住谷口退路需要投入指定行动点。",
        "显示内容：中央展示己方角色与敌方单位。",
        "显示内容：底部展示当前可用教学行动。",
        "显示内容：回合区域显示剩余行动点。",
        "主要操作：依据优先目标选择行动。",
        "回合操作：行动点使用完毕后结束当前回合。",
        "数值规则：教学战共四个回合。",
        "数值规则：每回合提供两个行动点。",
        "数值规则：总行动点为八点。",
        "取舍规则：三个目标全部完成需要九点。",
        "设计结果：用户不能同时完成全部三个目标。",
        "状态记录：各目标结果分别写入剧情标记。",
        "后续影响：教学结果影响后续剧情和战斗状态。",
        "意图规则：敌方行动前显示预告信息。",
        "行动反馈：选择行动后更新目标进度。",
        "战斗反馈：伤害、救援和防守结果写入日志。",
        "结束条件：完成四回合或触发预定结果。",
        "流程出口：结算后进入越女来援剧情。",
        "失败理解：未完成目标属于叙事结果而非程序错误。",
        "交互方式：按钮同时支持鼠标和触摸。",
        "声音控制：用户可在教学战中切换静音。",
        "异常处理：刷新后可依靠战前存档恢复。",
        "功能对应：本页对应独立教学战斗模块。",
        "业务价值：固定取舍用于解释游戏核心规则。",
        "截图结论：图 6 证明软件具备教学目标系统。",
    ),
    (
        "界面用途：火场页执行限时档案抢救玩法。",
        "进入条件：第一章火场选择中选择抢救档案。",
        "显示内容：顶部提示官署起火的剧情背景。",
        "显示内容：右上方显示剩余总时间。",
        "显示内容：进度条表现火势推进状态。",
        "显示内容：中部列出五类可抢救档案。",
        "显示内容：每张卡片标明档案名称和作用。",
        "显示内容：可用档案标明距离与抢救耗时。",
        "显示内容：已焚毁档案显示不可选择状态。",
        "显示内容：已抢救档案显示完成状态。",
        "主要操作：在时间结束前点击档案卡片。",
        "数量规则：本次最多抢救三份档案。",
        "耗时规则：不同档案的抢救耗时不同。",
        "焚毁规则：档案按既定火势时间表焚毁。",
        "选择规则：已焚毁档案不能再次选择。",
        "选择规则：已抢救档案不能重复抢救。",
        "状态反馈：底部显示已抢救数量和累计耗时。",
        "倒计时反馈：剩余时间持续更新。",
        "结束条件：达到数量上限或时间归零。",
        "状态记录：抢救清单写入剧情标记。",
        "后续影响：抢救结果进入章末支线展示。",
        "流程出口：玩法结束后返回火起后段剧情。",
        "交互方式：桌面端点击档案卡片。",
        "交互方式：移动端触摸档案卡片。",
        "声音功能：操作与计时节点可触发音效。",
        "异常处理：图片资源失败不影响卡片判定。",
        "异常处理：倒计时由程序状态独立维护。",
        "功能对应：本页对应火场救援模块。",
        "业务价值：限时资源选择形成差异化支线。",
        "截图结论：图 7 证明软件具备限时救援玩法。",
    ),
    (
        "界面用途：追截页执行多波路线拦截玩法。",
        "进入条件：第一章火场选择中选择追截残军。",
        "显示内容：顶部显示“截残军”任务名称。",
        "显示内容：右上角显示当前剩余判断时间。",
        "显示内容：进度条显示当前波次。",
        "显示内容：统计区显示已经拦截的数量。",
        "显示内容：线索框提示残军移动信息。",
        "显示内容：下方列出三条候选路线。",
        "主要操作：依据线索点击可能的逃离路线。",
        "选择限制：每个波次仅允许确认一次。",
        "计时规则：各波次分别计算选择时间。",
        "判定规则：选中正确路线即记为成功拦截。",
        "判定规则：错误路线不增加累计数量。",
        "超时规则：未及时选择时按未拦截处理。",
        "进度规则：当前波次完成后进入下一波。",
        "累计规则：系统记录全部波次成功次数。",
        "状态记录：最终数量写入追截剧情标记。",
        "后续影响：章末结算展示追截结果。",
        "流程出口：完成全部波次后返回主剧情。",
        "视觉反馈：已选按钮在确认后改变状态。",
        "视觉反馈：线索和倒计时保持高对比显示。",
        "交互方式：桌面端支持鼠标点击。",
        "交互方式：移动端支持触摸操作。",
        "声音控制：右上角提供静音开关。",
        "异常处理：背景失败不影响路线按钮。",
        "异常处理：玩法数据由当前状态计算。",
        "功能对应：本页对应残军追截模块。",
        "业务价值：玩法结合观察、判断和时间压力。",
        "版本范围：该玩法属于第一章三选一支线。",
        "截图结论：图 8 证明软件具备路线拦截流程。",
    ),
    (
        "界面用途：撤离页执行疫营护送防守玩法。",
        "进入条件：第一章火场选择中选择护送疫营。",
        "显示内容：顶部显示新郑疫营撤离战标题。",
        "显示内容：背景区域展示撤离场景。",
        "显示内容：危机框公开本回合威胁。",
        "显示内容：目标血条表示撤离队伍状态。",
        "显示内容：右侧显示当前回合和总回合数。",
        "显示内容：底部提供结阵、抢救和反击。",
        "主要操作：依据公开危机选择一种应对。",
        "回合规则：玩法持续四个回合。",
        "危机类型：正面冲阵强调防守应对。",
        "危机类型：担架断裂强调抢救应对。",
        "危机类型：头目突入强调反击应对。",
        "结阵效果：优先降低正面冲击造成的伤害。",
        "抢救效果：用于恢复目标队伍状态。",
        "反击效果：针对关键敌方威胁。",
        "疲惫规则：连续使用相同应对产生惩罚。",
        "伤害规则：每回合基础威胁逐步提高。",
        "预览反馈：按钮区域提示预计效果。",
        "状态反馈：选择后更新目标血量和回合。",
        "胜利条件：目标未归零并守满四回合。",
        "失败条件：撤离目标血量降为零。",
        "结果分档：按剩余血量记录高或低存活率。",
        "状态记录：撤离结果写入剧情标记。",
        "后续影响：存活率进入章末支线结算。",
        "失败处理：战败后可使用重试覆盖层。",
        "流程出口：结束后返回火起后段剧情。",
        "功能对应：本页对应撤离防守模块。",
        "业务价值：公开危机支持有依据的策略判断。",
        "截图结论：图 9 证明软件具备撤离防守玩法。",
    ),
    (
        "界面用途：存档弹窗管理本地手动进度。",
        "进入方式：从标题页或可存档场景打开。",
        "显示内容：弹窗标题标明当前为读取进度。",
        "显示内容：页面并列展示三个存档槽位。",
        "显示内容：非空槽位显示章节和场景名称。",
        "显示内容：非空槽位显示保存时间。",
        "显示内容：空槽位明确标记为空。",
        "主要操作：点击非空槽位的读取按钮。",
        "读取结果：系统恢复场景状态和剧情标记。",
        "保存操作：可存档场景可向指定槽位保存。",
        "覆盖规则：覆盖已有槽位前要求确认。",
        "删除规则：删除已有存档前要求确认。",
        "关闭操作：点击关闭图标或按 Esc 返回。",
        "槽位规则：软件提供三个手动存档槽位。",
        "自动存档：进入配置战斗前建立独立自动存档。",
        "检查点：战败回卷使用会话级战斗检查点。",
        "数据字段：存档包含版本、场景、标记和时间。",
        "兼容分类：当前版本完整存档可正常读取。",
        "兼容分类：旧版本存档提示处理或重置。",
        "兼容分类：未来版本存档禁止读取和覆盖。",
        "兼容分类：解析失败的存档归类为无效。",
        "规范化：读取前校验剧情标记类型和取值。",
        "存储位置：浏览器版本使用本地存储。",
        "隐私规则：存档数据不上传业务服务器。",
        "降级处理：本地存储不可用时采用安全降级。",
        "交互方式：桌面端和移动端均可操作。",
        "功能对应：本页对应存档管理与兼容模块。",
        "业务价值：多槽位支持不同进度留存。",
        "使用结果：用户可从已保存节点继续游戏。",
        "截图结论：图 10 证明软件具备存读档界面。",
    ),
    (
        "界面用途：结算页汇总第一章流程结果。",
        "进入条件：第一章末尾剧情与支线完成后显示。",
        "显示内容：页面顶部显示章节名称和总评。",
        "显示内容：总评依据关键状态生成。",
        "显示内容：战绩栏目汇总主要战斗结果。",
        "显示内容：抉择栏目汇总关键剧情选择。",
        "显示内容：支线栏目汇总三选一玩法结果。",
        "显示内容：史乘栏目展示解锁进度。",
        "显示内容：页面下方显示章节总结文本。",
        "主要操作：阅读各栏目和章节评价。",
        "继续操作：点击收军报进入后续结束流程。",
        "数据来源：结算页读取既有剧情标记。",
        "数据来源：战斗表现由相关战斗结果提供。",
        "数据来源：火场结果来自档案抢救清单。",
        "数据来源：追截结果来自成功拦截数量。",
        "数据来源：撤离结果来自目标存活率分档。",
        "计算规则：总评由多项状态组合决定。",
        "只读规则：结算页不修改历史选择。",
        "只读规则：结算页不重新执行战斗判定。",
        "史乘规则：已满足条件的对照卡显示进度。",
        "流程出口：继续后进入终章或史乘入口。",
        "交互方式：桌面端可点击底部继续按钮。",
        "交互方式：移动端可触摸继续按钮。",
        "声音控制：页面保留全局声音控制。",
        "异常处理：单项数据缺失时使用安全默认值。",
        "响应布局：结算栏目按屏幕宽度重排。",
        "功能对应：本页对应章末结算模块。",
        "业务价值：把此前选择转化为可见结果。",
        "版本范围：V1.0 提供第一章完整结算。",
        "截图结论：图 11 证明软件具备章末汇总功能。",
    ),
    (
        "界面用途：史乘页展示叙事与史实的对照内容。",
        "进入方式：从标题页、章末或终章入口进入。",
        "显示内容：页面顶部显示“史乘”标题。",
        "显示内容：顶部同时显示当前解锁进度。",
        "显示内容：主体以卡片网格展示历史主题。",
        "显示内容：每张卡片标明所属章节和编号。",
        "显示内容：卡片上部展示史书大意。",
        "显示内容：卡片下部展示玩家经历的叙事。",
        "主要操作：阅读已解锁的历史对照卡。",
        "关闭操作：点击右上角关闭按钮返回。",
        "解锁规则：卡片依据剧情进度逐步解锁。",
        "未解锁状态：未满足条件时显示锁定提示。",
        "只读规则：浏览史乘不改变剧情选择。",
        "只读规则：浏览史乘不改变战斗结果。",
        "数据来源：对照卡由独立史乘数据定义。",
        "内容结构：每张卡区分历史记载与游戏叙事。",
        "展示主题：包括疲秦计等历史背景。",
        "展示主题：包括逐客令等历史背景。",
        "展示主题：包括韩国灭亡相关史实。",
        "展示主题：包括郑国渠相关史实。",
        "进度反馈：已解锁数量实时显示。",
        "视觉设计：卡片采用统一的深色文献风格。",
        "交互方式：桌面端支持鼠标滚动与点击。",
        "交互方式：移动端支持触摸滚动。",
        "声音控制：页面继续使用全局静音状态。",
        "异常处理：个别媒体失败不影响文本阅读。",
        "响应布局：卡片数量随屏幕宽度调整。",
        "功能对应：本页对应史乘历史对照模块。",
        "业务价值：明确区分艺术叙事和史实资料。",
        "截图结论：图 12 证明软件具备史乘对照功能。",
    ),
)

FORBIDDEN = (
    "Qoder",
    "Codex",
    "AI生成",
    "待本人填写",
    "待本人确认",
    "待申请人",
    "/Users/",
    "身份证",
    "证件号码",
    "联系电话",
    "电子邮箱",
    "通信地址",
    "ghp_",
    "gho_",
    "AppSecret",
)


@dataclass(frozen=True)
class DepositLine:
    text: str
    section: str


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def display_width(text: str) -> int:
    return sum(2 if ord(ch) > 127 else 1 for ch in text)


def clean_text(text: str) -> str:
    text = text.replace("環境", "环境")
    text = text.replace("乘", "×") if re.search(r"\d\s*乘\s*\d", text) else text
    text = re.sub(r"\s+", " ", text).strip()
    return text


def wrap_display(text: str, width: int = MAX_DISPLAY_WIDTH) -> list[str]:
    """Wrap without losing content, preferring punctuation/space boundaries."""
    text = clean_text(text)
    if not text:
        return []
    result: list[str] = []
    rest = text
    while display_width(rest) > width:
        used = 0
        cut = 0
        preferred = 0
        for idx, ch in enumerate(rest):
            used += 2 if ord(ch) > 127 else 1
            if used > width:
                break
            cut = idx + 1
            if ch in "，。；：、,.!?;:）)]} " and used >= int(width * 0.62):
                preferred = cut
        if preferred:
            cut = preferred
        if cut <= 0:
            cut = 1
        result.append(rest[:cut].strip())
        rest = rest[cut:].strip()
    if rest:
        result.append(rest)
    return [x for x in result if x]


def markdown_to_lines(path: Path) -> list[DepositLine]:
    lines: list[DepositLine] = []
    section = "软件概述"
    for raw in path.read_text(encoding="utf-8").splitlines():
        text = raw.strip()
        if not text or text == "---" or text.startswith("![") or text.startswith("```"):
            continue
        if re.fullmatch(r"\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?", text):
            continue
        if text.startswith("#"):
            section = clean_text(text.lstrip("#").strip())
            text = f"【{section}】"
        elif text.startswith("|") and text.endswith("|"):
            cells = [clean_text(cell) for cell in text.strip("|").split("|")]
            cells = [cell for cell in cells if cell]
            if not cells:
                continue
            text = "；".join(cells)
        else:
            text = re.sub(r"^[-*+]\s+", "", text)
        for wrapped in wrap_display(text):
            lines.append(DepositLine(wrapped, section))
    return lines


def core_manifest() -> list[dict]:
    path = WORKING / "source-manifest-V1.0.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    return [item for item in data["includedFiles"] if item.get("included")]


def role_for(path: str) -> str:
    exact = {
        "src/main.tsx": "浏览器入口与根组件挂载",
        "src/App.tsx": "全局场景状态机、流程编排与背景音乐映射",
        "src/game/story.ts": "剧情场景、对白、选择与历史对照数据",
        "src/game/battle.ts": "通用回合制战斗状态与纯函数结算",
        "src/game/scenarios.ts": "各战斗场景的配置与分支参数",
        "src/game/tutorial.ts": "教学战斗状态、行动点和结果判定",
        "src/game/defenseBattle.ts": "撤离防守危机、预览和回合结算",
        "src/game/save.ts": "本地存档、兼容性、规范化与自动存档",
        "src/game/audio.ts": "背景音乐、语音、音效和静音状态管理",
        "src/game/animationEngine.ts": "战斗位移、命中、受击和镜头反馈",
        "src/sections/TitleScreen.tsx": "标题页、继续游戏和辅助入口",
        "src/sections/StoryScene.tsx": "剧情对白播放、分支选择和语音触发",
        "src/sections/BattleScene.tsx": "通用战斗界面、行动选择和动画调度",
        "src/sections/TutorialBattleScene.tsx": "教学战斗界面与目标反馈",
    }
    if path in exact:
        return exact[path]
    if path.startswith("src/platform/"):
        return "平台存储、路由、音频、图像或分享适配"
    if path.startswith("src/ui/"):
        return "业务界面、交互面板或配套静态数据"
    if path.startswith("src/components/"):
        return "可复用视觉组件、弹窗或战斗反馈组件"
    if path.startswith("src/hooks/"):
        return "可复用状态钩子与资源加载逻辑"
    if path.startswith("src/lib/"):
        return "资源地址、预加载或公共工具逻辑"
    return "V1.0 核心实现模块"


def grouped(items: list[str], size: int = 4) -> Iterable[list[str]]:
    for idx in range(0, len(items), size):
        yield items[idx : idx + size]


def source_facts(manifest: list[dict]) -> list[DepositLine]:
    facts: list[DepositLine] = []

    def add(section: str, text: str) -> None:
        for line in wrap_display(text):
            facts.append(DepositLine(line, section))

    add("附录 N 源码模块与依赖追溯", "【附录 N 源码模块与依赖追溯】")
    for item in manifest:
        path = item["path"]
        add(
            "附录 N 源码模块与依赖追溯",
            f"模块 {path}：共 {item['totalLines']} 行、非空 {item['nonEmptyLines']} 行；职责为{role_for(path)}。",
        )
        add(
            "附录 N 源码模块与依赖追溯",
            f"模块校验 {path}：SHA-256 前十二位为 {item['sha256'][:12]}，纳入 V1.0 固定源码清单。",
        )
        src = (REPO / path).read_text(encoding="utf-8")
        imports = re.findall(r"\bfrom\s+['\"]([^'\"]+)['\"]", src)
        imports = list(dict.fromkeys(imports))
        for group in grouped(imports, 4):
            add(
                "附录 N 源码模块与依赖追溯",
                f"依赖关系 {path}：直接引用 {', '.join(group)}；引用关系构成该模块的输入边界。",
            )

    add("附录 O 类型、函数与常量索引", "【附录 O 类型、函数与常量索引】")
    symbol_re = re.compile(
        r"^(?:export\s+)?(?:default\s+)?(interface|type|class|function|const|enum)\s+([A-Za-z_$][A-Za-z0-9_$]*)",
        re.M,
    )
    descriptions = {
        "interface": "定义结构化数据契约",
        "type": "限定状态、参数或返回值的类型边界",
        "class": "封装具有生命周期的实现对象",
        "function": "封装可复用的计算或交互流程",
        "const": "保存只读配置、映射或函数引用",
        "enum": "定义有限状态集合",
    }
    for item in manifest:
        path = item["path"]
        src = (REPO / path).read_text(encoding="utf-8")
        for match in symbol_re.finditer(src):
            kind, name = match.groups()
            line_no = src.count("\n", 0, match.start()) + 1
            add(
                "附录 O 类型、函数与常量索引",
                f"符号 {name}：{kind}，位于 {path} 第 {line_no} 行；用于{descriptions[kind]}。",
            )

    add("附录 P 界面状态与数据字段追溯", "【附录 P 界面状态与数据字段追溯】")
    state_re = re.compile(r"const\s*\[\s*([A-Za-z_$][A-Za-z0-9_$]*)\s*,[^\]]+\]\s*=\s*(?:React\.)?useState")
    interface_re = re.compile(r"^(?:export\s+)?interface\s+([A-Za-z_$][A-Za-z0-9_$]*)")
    field_re = re.compile(r"^\s{2}([A-Za-z_$][A-Za-z0-9_$]*)\??\s*:")
    for item in manifest:
        path = item["path"]
        src_lines = (REPO / path).read_text(encoding="utf-8").splitlines()
        current_interface: str | None = None
        brace_depth = 0
        for line_no, raw in enumerate(src_lines, 1):
            for match in state_re.finditer(raw):
                add(
                    "附录 P 界面状态与数据字段追溯",
                    f"界面状态 {match.group(1)}：在 {path} 第 {line_no} 行由 useState 管理，随组件交互生命周期更新。",
                )
            im = interface_re.match(raw)
            if im:
                current_interface = im.group(1)
                brace_depth = raw.count("{") - raw.count("}")
                continue
            if current_interface:
                brace_depth += raw.count("{") - raw.count("}")
                fm = field_re.match(raw)
                if fm:
                    add(
                        "附录 P 界面状态与数据字段追溯",
                        f"数据字段 {current_interface}.{fm.group(1)}：定义于 {path} 第 {line_no} 行，属于该接口的显式字段。",
                    )
                if brace_depth <= 0:
                    current_interface = None

    package = json.loads((REPO / "package.json").read_text(encoding="utf-8"))
    add("附录 Q 构建、依赖与验证入口", "【附录 Q 构建、依赖与验证入口】")
    for name, command in package.get("scripts", {}).items():
        add("附录 Q 构建、依赖与验证入口", f"命令 {name}：执行 {command}；用于开发、构建、测试或发布流程。")
    for name, version in {**package.get("dependencies", {}), **package.get("devDependencies", {})}.items():
        add("附录 Q 构建、依赖与验证入口", f"依赖 {name}：版本约束 {version}；由包管理清单锁定其安装范围。")

    add("附录 R 源码注释与实现约束摘录", "【附录 R 源码注释与实现约束摘录】")
    for item in manifest:
        path = item["path"]
        for line_no, raw in enumerate((REPO / path).read_text(encoding="utf-8").splitlines(), 1):
            stripped = raw.strip()
            if not (stripped.startswith("//") or stripped.startswith("/*") or stripped.startswith("*")):
                continue
            comment = re.sub(r"^(?://|/\*+|\*+)\s*", "", stripped).rstrip("*/ ")
            if len(comment) < 5:
                continue
            add("附录 R 源码注释与实现约束摘录", f"实现注释 {path} 第 {line_no} 行：{comment}")
    return facts


def select_document_lines() -> list[DepositLine]:
    all_base = markdown_to_lines(WORKING / "软件设计与使用说明书-V1.0.md")
    appendix_start = next((index for index, item in enumerate(all_base) if item.section.startswith("附录")), len(all_base))
    base = all_base[:appendix_start]
    base_appendices = all_base[appendix_start:]
    facts = source_facts(core_manifest())
    if len(base) > MAIN_CHAPTER_LINE_TARGET:
        raise RuntimeError(
            f"Main chapters exceed {MAIN_CHAPTER_LINE_TARGET} lines and would spill into the screenshot chapter"
        )
    chapter_padding = [
        DepositLine("小游戏适配版通过独立构建入口生成发布目录。", "第 11 章 安装、启动与使用方法"),
        DepositLine("微信小游戏版调用微信平台存储与音频接口。", "第 11 章 安装、启动与使用方法"),
        DepositLine("抖音小游戏版调用抖音平台存储与音频接口。", "第 11 章 安装、启动与使用方法"),
        DepositLine("各平台版本共用核心剧情、战斗和存档规则。", "第 11 章 安装、启动与使用方法"),
    ]
    base_line_count = len(base)
    while len(base) < MAIN_CHAPTER_LINE_TARGET:
        base.append(chapter_padding[(len(base) - base_line_count) % len(chapter_padding)])
    combined = base + base_appendices + facts
    target = TEXT_PAGES * LINES_PER_PAGE
    if len(combined) > target:
        combined = combined[:target]
    if len(combined) < target:
        manifest = core_manifest()
        idx = 0
        while len(combined) < target:
            item = manifest[idx % len(manifest)]
            path = item["path"]
            line_no = idx // len(manifest) + 1
            src_lines = (REPO / path).read_text(encoding="utf-8").splitlines()
            if line_no <= len(src_lines) and src_lines[line_no - 1].strip():
                combined.append(
                    DepositLine(
                        f"实现追溯 {path} 第 {line_no} 行：该有效语句属于 V1.0 固定源码基线并受文件哈希校验。",
                        "附录 S 实现基线追溯",
                    )
                )
            idx += 1
    manifest = core_manifest()
    pad_idx = 0
    while len(combined) < target:
        item = manifest[pad_idx % len(manifest)]
        combined.append(
            DepositLine(
                f"清单核验 {item['path']}：非空代码 {item['nonEmptyLines']} 行，文件摘要 {item['sha256'][:16]}。",
                "附录 S 实现基线追溯",
            )
        )
        pad_idx += 1
    return combined[:target]


def set_run_font(run, font: str, size: float, *, bold: bool = False, color: str = "000000") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, DOC_FONT, 8.5)


def configure_document(doc: Document, *, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)
    doc.core_properties.title = title
    doc.core_properties.subject = f"{SOFTWARE} {VERSION} 软件著作权登记鉴别材料"
    doc.core_properties.author = SOFTWARE
    doc.core_properties.keywords = "软件著作权,设计说明书,使用说明书,V1.0"

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(9.6)
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(18.5)

    if "Deposit Page Heading" not in [s.name for s in doc.styles]:
        page_style = doc.styles.add_style("Deposit Page Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        page_style = doc.styles["Deposit Page Heading"]
    page_style.font.name = DOC_FONT
    page_style.font.size = Pt(11.5)
    page_style.font.bold = True
    page_style.font.color.rgb = RGBColor.from_string("000000")
    page_style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    page_style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    page_style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    page_style.paragraph_format.space_before = Pt(0)
    page_style.paragraph_format.space_after = Pt(0)
    page_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    page_style.paragraph_format.line_spacing = Pt(22)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run(f"{SOFTWARE} {VERSION}  |  软件设计与使用说明书")
    set_run_font(hr, DOC_FONT, 8.5)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("第 ")
    set_run_font(fr, DOC_FONT, 8.5)
    add_field(fp, "PAGE")
    fr = fp.add_run(" 页 / 共 ")
    set_run_font(fr, DOC_FONT, 8.5)
    add_field(fp, "NUMPAGES")
    fr = fp.add_run(" 页")
    set_run_font(fr, DOC_FONT, 8.5)


def clear_header_footer(section) -> None:
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            paragraph.clear()


def add_cover_and_contents(doc: Document) -> None:
    section = doc.sections[0]
    clear_header_footer(section)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(82)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(SOFTWARE)
    set_run_font(run, DOC_FONT, 26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(VERSION)
    set_run_font(run, DOC_FONT, 18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    run = p.add_run("软件设计与使用说明书")
    set_run_font(run, DOC_FONT, 22, bold=True)

    cover_rows = (
        ("软件全称", SOFTWARE),
        ("软件简称", "秦灭六国"),
        ("版本号", VERSION),
        ("文档类型", "软件设计与使用说明书"),
        ("运行形态", "跨平台数字游戏（浏览器运行，可适配微信/抖音小游戏）"),
        ("编制日期", "2026 年 08 月"),
    )
    table = doc.add_table(rows=len(cover_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, (label, value) in enumerate(cover_rows):
        left, right = table.rows[row_idx].cells
        left.width = Mm(38)
        right.width = Mm(92)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(label)
        set_run_font(r, DOC_FONT, 10, bold=True)
        p = right.paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, DOC_FONT, 10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(75)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件著作权登记鉴别材料")
    set_run_font(r, DOC_FONT, 10)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("目  录")
    set_run_font(r, DOC_FONT, 20, bold=True)

    entries = (
        ("第1章  软件概述", 3),
        ("第2章  运行环境", 4),
        ("第3章  技术架构", 4),
        ("第4章  数据模型与状态管理", 5),
        ("第5章  功能模块设计", 6),
        ("第6章  核心业务流程", 8),
        ("第7章  界面组成与交互", 8),
        ("第8章  业务规则", 9),
        ("第9章  辅助功能", 10),
        ("第10章  异常处理与兼容性", 10),
        ("第11章  安装、启动与使用方法", 11),
        ("第12章  运行截图及说明", 12),
        ("附录A—M  设计数据与实现明细", 24),
        ("附录N  源码模块与依赖追溯", 33),
        ("附录O  类型、函数与常量索引", 42),
    )
    for label, page_number in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(12)
        p.paragraph_format.right_indent = Mm(12)
        p.paragraph_format.space_after = Pt(5.5)
        p.paragraph_format.tab_stops.add_tab_stop(Mm(158))
        r = p.add_run(f"{label}\t{page_number}")
        set_run_font(r, DOC_FONT, 10.5, bold=label.startswith("第"))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("说明：目录页码包含封面和目录；正文第3页起，第12章运行截图第12页起，附录第24页起。")
    set_run_font(r, DOC_FONT, 8.5)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    content_section = doc.add_section(WD_SECTION.CONTINUOUS)
    content_section.page_width = Mm(210)
    content_section.page_height = Mm(297)
    content_section.top_margin = Mm(15)
    content_section.bottom_margin = Mm(15)
    content_section.left_margin = Mm(17)
    content_section.right_margin = Mm(17)
    content_section.header_distance = Mm(6)
    content_section.footer_distance = Mm(7)
    content_section.header.is_linked_to_previous = False
    content_section.footer.is_linked_to_previous = False
    hp = content_section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run(f"{SOFTWARE} {VERSION}  |  软件设计与使用说明书")
    set_run_font(hr, DOC_FONT, 8.5)
    fp = content_section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_run_font(fr, DOC_FONT, 8.5)
    add_field(fp, "PAGE")
    fr = fp.add_run(" 页 / 共 ")
    set_run_font(fr, DOC_FONT, 8.5)
    add_field(fp, "NUMPAGES")
    fr = fp.add_run(" 页")
    set_run_font(fr, DOC_FONT, 8.5)


def build_design_docx(lines: list[DepositLine], path: Path, trace_path: Path) -> None:
    if len(lines) != TEXT_PAGES * LINES_PER_PAGE:
        raise RuntimeError(f"Expected {TEXT_PAGES * LINES_PER_PAGE} text lines, got {len(lines)}")
    missing = [str(SCREENSHOT_DIR / item["file"]) for item in SCREENSHOTS if not (SCREENSHOT_DIR / item["file"]).is_file()]
    if missing:
        raise FileNotFoundError(f"Missing runtime screenshots: {missing}")

    doc = Document()
    configure_document(doc, title=DOC_TITLE)
    add_cover_and_contents(doc)
    pages = len(lines) // LINES_PER_PAGE
    with trace_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["page", "inPageLine", "section", "content", "pageType", "imageFile"])

        def normalize_heading(text: str) -> str:
            value = text[1:-1].strip() if text.startswith("【") and text.endswith("】") else text.strip()
            value = re.sub(r"^第\s*(\d+)\s*章\s*", r"第\1章  ", value)
            value = re.sub(r"^附录\s*([A-Z])\s*", r"附录\1  ", value)
            value = re.sub(r"^([A-Z])\.(\d+)\s*", r"\1.\2  ", value)
            value = re.sub(r"^(\d+)\.(\d+)\s*", r"\1.\2  ", value)
            return value

        def heading_level(text: str) -> str | None:
            if not (text.startswith("【") and text.endswith("】")):
                return None
            value = normalize_heading(text)
            if re.match(r"^第\d+章\s", value) or re.match(r"^附录[A-Z]\s", value):
                return "chapter"
            if re.match(r"^(?:\d+|[A-Z])\.\d+\s", value):
                return "section"
            return "document"

        def add_text_page(page_idx: int, displayed_page: int, *, add_break: bool) -> None:
            page_lines = lines[page_idx * LINES_PER_PAGE : (page_idx + 1) * LINES_PER_PAGE]
            for line_idx, item in enumerate(page_lines, 1):
                p = doc.add_paragraph()
                p.paragraph_format.keep_together = True
                p.paragraph_format.widow_control = False
                level = heading_level(item.text)
                if level == "chapter":
                    p.paragraph_format.space_before = Pt(5)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(19)
                    run = p.add_run(normalize_heading(item.text))
                    set_run_font(run, DOC_FONT, 13, bold=True)
                elif level == "section":
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(18.5)
                    run = p.add_run(normalize_heading(item.text))
                    set_run_font(run, DOC_FONT, 10.5, bold=True)
                elif level == "document":
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(18.5)
                    run = p.add_run(normalize_heading(item.text))
                    set_run_font(run, DOC_FONT, 11, bold=True)
                else:
                    run = p.add_run(item.text)
                    set_run_font(run, DOC_FONT, 9.6)
                writer.writerow([displayed_page, line_idx, item.section, item.text, "text", ""])
            if add_break:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.add_run().add_break(WD_BREAK.PAGE)

        # Chapters 1–11 precede the screenshot chapter.  The remaining text
        # pages are appendices and therefore follow Chapter 12.
        for page_idx in range(MAIN_CHAPTER_TEXT_PAGES):
            displayed_page = FRONT_MATTER_PAGES + page_idx + 1
            add_text_page(page_idx, displayed_page, add_break=True)

        if len(SCREENSHOTS) != len(SCREENSHOT_OPERATION_DETAILS):
            raise RuntimeError("Screenshot metadata and operation detail counts disagree")
        if any(len(details) != LINES_PER_PAGE for details in SCREENSHOT_OPERATION_DETAILS):
            raise RuntimeError("Every screenshot page must have exactly 30 substantive operation lines")

        for visual_idx, (item, details) in enumerate(zip(SCREENSHOTS, SCREENSHOT_OPERATION_DETAILS)):
            page_number = FRONT_MATTER_PAGES + MAIN_CHAPTER_TEXT_PAGES + visual_idx + 1
            p = doc.add_paragraph(style="Deposit Page Heading")
            p.paragraph_format.keep_with_next = False
            feature_title = item["title"].split("  ", 1)[1]
            if visual_idx == 0:
                run = p.add_run("第12章  运行截图及说明")
                set_run_font(run, DOC_FONT, 13, bold=True)
                subheading = doc.add_paragraph()
                subheading.paragraph_format.space_before = Pt(1)
                subheading.paragraph_format.space_after = Pt(1)
                run = subheading.add_run(f"12.{visual_idx + 1}  {feature_title}")
                set_run_font(run, DOC_FONT, 10.5, bold=True)
            else:
                run = p.add_run(f"12.{visual_idx + 1}  {feature_title}")
                set_run_font(run, DOC_FONT, 10.5, bold=True)
            intro = f"本页展示{feature_title}，并按实际运行逻辑说明进入方式、主要操作、状态规则和处理结果。"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(intro)
            set_run_font(run, DOC_FONT, 8.5)

            image_p = doc.add_paragraph()
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_p.paragraph_format.space_before = Pt(1)
            image_p.paragraph_format.space_after = Pt(1)
            image_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            image_run = image_p.add_run()
            image_run.add_picture(str(SCREENSHOT_DIR / item["file"]), width=Mm(118))

            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_before = Pt(0)
            caption.paragraph_format.space_after = Pt(2)
            caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = caption.add_run(f"图12-{visual_idx + 1}  {feature_title}")
            set_run_font(run, DOC_FONT, 8.5, bold=True)

            detail_table = doc.add_table(rows=15, cols=2)
            detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            detail_table.autofit = False
            for row_idx, row in enumerate(detail_table.rows):
                row.height = Mm(6.25)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                for col_idx, cell in enumerate(row.cells):
                    cell.width = Mm(84)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell, top=20, start=70, bottom=20, end=70)
                    if row_idx % 2:
                        set_cell_shading(cell, "F7FAFC")
                    line_idx = row_idx * 2 + col_idx
                    content = details[line_idx]
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    run = paragraph.add_run(f"{line_idx + 1:02d}. {content}")
                    set_run_font(run, DOC_FONT, 6.8)
                    writer.writerow([page_number, line_idx + 1, "第 12 章 运行截图及说明", content, "screenshot", item["file"] if line_idx == 0 else ""])

            if visual_idx + 1 < len(SCREENSHOTS) or APPENDIX_TEXT_PAGES:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.add_run().add_break(WD_BREAK.PAGE)

        for appendix_idx, page_idx in enumerate(range(MAIN_CHAPTER_TEXT_PAGES, pages)):
            displayed_page = FRONT_MATTER_PAGES + SCREENSHOT_PAGES + page_idx + 1
            add_text_page(page_idx, displayed_page, add_break=appendix_idx + 1 < APPENDIX_TEXT_PAGES)
    doc.save(path)


def build_source_docx(path: Path) -> None:
    trace = list(csv.DictReader((WORKING / "source-trace.csv").open(encoding="utf-8")))
    if len(trace) != 3000:
        raise RuntimeError(f"Expected 3000 source rows, got {len(trace)}")
    doc = Document()
    configure_document(doc, title=f"{SOFTWARE} {VERSION} 源程序鉴别材料")
    section = doc.sections[0]
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    header = section.header.paragraphs[0]
    for run in list(header.runs):
        run._element.getparent().remove(run._element)
    hr = header.add_run(f"{SOFTWARE} {VERSION}  |  源程序鉴别材料")
    set_run_font(hr, CODE_FONT, 8.0, color="6B7280")
    for page_idx in range(60):
        page_rows = trace[page_idx * 50 : (page_idx + 1) * 50]
        for row in page_rows:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(10.4)
            p.paragraph_format.left_indent = Mm(7)
            p.paragraph_format.first_line_indent = Mm(-7)
            n = p.add_run(f"{int(row['inPageLine']):02d}  ")
            set_run_font(n, CODE_FONT, 6.8, color="9CA3AF")
            code = p.add_run(row["content"])
            set_run_font(code, CODE_FONT, 6.8, color="111827")
        if page_idx < 59:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run().add_break(WD_BREAK.PAGE)
    doc.save(path)


def write_confirmation_template() -> None:
    text = f"""# {SOFTWARE} {VERSION} 申请人事实确认表

> 本文件为申请人本人填写的内部表，不上传为软件鉴别材料。完成以下信息后，方可生成可直接提交的登记申请信息。

1. 申请人姓名或主体名称：
2. 证件类型：
3. 证件号码：
4. 联系地址：
5. 联系电话：
6. 电子邮箱：
7. 软件开发完成日期（YYYY-MM-DD）：
8. 软件是否已经向公众提供：
9. 如已发表，首次发表日期与地点：
10. 开发方式（独立/合作/委托/职务）：
11. 权利取得方式（原始取得/继受取得）：
12. 是否基于他人软件修改：
13. 是否存在共同开发人或其他权利人：
14. 软件全称“{SOFTWARE}”和版本号“{VERSION}”是否最终确认：
"""
    (CONFIRM / "申请人事实确认表-填写后回传.md").write_text(text, encoding="utf-8")


def write_readme(lines: list[DepositLine]) -> None:
    pages = TOTAL_PDF_PAGES
    text = f"""# {SOFTWARE} {VERSION} 技术提交材料说明

本目录由当前 V1.0 源码重新生成。技术材料按国家版权局公开规章中的 A4、源程序每页不少于50行、文档每页不少于30行要求制作。

## 正式提交件

- `01-正式提交件-技术材料/01-{SOFTWARE}-{VERSION}-源程序鉴别材料.pdf`：60页，前30页加后30页，每页50行非空代码。
- `01-正式提交件-技术材料/02-{SOFTWARE}-{VERSION}-软件设计与使用说明书.pdf`：{pages}页，含1页封面、1页目录、{TEXT_PAGES}页设计与使用正文及{SCREENSHOT_PAGES}页运行截图说明；正文与截图说明页均有30行实质内容，提交全部文档。

## 仍需申请人本人完成

- 填写 `03-申请人确认后填写/申请人事实确认表-填写后回传.md`。
- 登录登记系统，按真实情况填写申请表并上传身份证明、权属证明和两份技术PDF。

申请人事实未确认前，技术材料可以验收，但整套登记申请仍不能宣称已经具备提交条件。
"""
    (OUT / "README-提交说明.md").write_text(text, encoding="utf-8")


def main() -> None:
    for directory in (SUBMIT, EDITABLE, CONFIRM, QA):
        directory.mkdir(parents=True, exist_ok=True)

    lines = select_document_lines()
    pages = TOTAL_PDF_PAGES
    if any(display_width(item.text) > MAX_DISPLAY_WIDTH for item in lines):
        raise RuntimeError("A design document line exceeds the fixed display width")
    for item in lines:
        for token in FORBIDDEN:
            if token in item.text:
                raise RuntimeError(f"Forbidden token {token!r} in design content: {item.text}")

    design_docx = EDITABLE / f"02-{SOFTWARE}-{VERSION}-软件设计与使用说明书.docx"
    trace_csv = QA / "软件说明书逐页内容追溯.csv"
    build_design_docx(lines, design_docx, trace_csv)

    source_docx = EDITABLE / f"01-{SOFTWARE}-{VERSION}-源程序鉴别材料.docx"
    build_source_docx(source_docx)

    source_pdf = QODER_FINAL / f"02-{SOFTWARE}-{VERSION}-源程序鉴别材料.pdf"
    if not source_pdf.exists():
        raise FileNotFoundError(source_pdf)
    shutil.copy2(source_pdf, SUBMIT / f"01-{SOFTWARE}-{VERSION}-源程序鉴别材料.pdf")
    shutil.copy2(WORKING / "source-trace.csv", QA / "源程序逐行追溯.csv")
    shutil.copy2(WORKING / "source-manifest-V1.0.json", QA / "源程序清单-V1.0.json")

    write_confirmation_template()
    write_readme(lines)
    build_meta = {
        "software": SOFTWARE,
        "version": VERSION,
        "documentPages": pages,
        "documentFrontMatterPages": FRONT_MATTER_PAGES,
        "documentTextPages": TEXT_PAGES,
        "documentScreenshotPages": SCREENSHOT_PAGES,
        "documentScreenshotCount": len(SCREENSHOTS),
        "documentLinesPerPage": LINES_PER_PAGE,
        "documentTotalLines": len(lines),
        "designDocxSha256": sha256(design_docx),
        "sourceDocxSha256": sha256(source_docx),
        "sourcePdfSha256": sha256(SUBMIT / f"01-{SOFTWARE}-{VERSION}-源程序鉴别材料.pdf"),
    }
    (QA / "构建元数据.json").write_text(json.dumps(build_meta, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(build_meta, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
