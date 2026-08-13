#!/usr/bin/env python3
"""Repair the page-number footer in the editable source deposit DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


FONT = "Arial Unicode MS"


def set_run_font(run, size: float = 8.5) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run)


def repair(path: Path) -> None:
    document = Document(path)
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run("第 ")
        set_run_font(run)
        add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页 / 共 ")
        set_run_font(run)
        add_field(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        set_run_font(run)
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()
    repair(args.docx)


if __name__ == "__main__":
    main()
