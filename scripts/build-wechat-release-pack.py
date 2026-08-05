#!/usr/bin/env python3
"""Build the consolidated WeChat release-material DOCX from Markdown sources."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACK_DIR = ROOT / "docs" / "release" / "微信上线待做事项"
OUTPUT = PACK_DIR / "秦灭六国_微信小游戏上架材料包_V1.0.docx"

SOURCE_FILES = [
    "README.md",
    "01-平台基础信息与提审文案.md",
    "02-小程序备案信息预填表.md",
    "03-隐私政策与个人信息清单.md",
    "04-用户服务协议.md",
    "05-适龄提示与未成年人保护.md",
    "06-游戏自审自查报告.md",
    "07-版权承诺与素材证据清单.md",
    "08-审核体验说明与截图计划.md",
    "09-上线操作清单与本人待办.md",
    "10-项目提审事实审计.md",
]

# compact_reference_guide preset, with a named CJK font override.
FONT_LATIN = "Calibri"
FONT_CJK = "PingFang SC"
INK = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 99, 112)
LIGHT_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    name = "Menlo" if mono else FONT_LATIN
    east = "Menlo" if mono else FONT_CJK
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)
    paragraph._p.append(r)
    end = paragraph.add_run(" 页")
    set_run_font(end, size=9, color=MUTED)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def compute_widths(rows):
    col_count = len(rows[0])
    weights = []
    for idx in range(col_count):
        max_len = max(len(row[idx]) if idx < len(row) else 0 for row in rows)
        weights.append(max(8, min(max_len, 48)))
    total = sum(weights)
    widths = [max(900, round(TABLE_WIDTH_DXA * weight / total)) for weight in weights]
    diff = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += diff
    return widths


def add_inline_markdown(paragraph, text, default_color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, color=default_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, color=INK, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, color=default_color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Title": (28, INK, 0, 8),
        "Subtitle": (13, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, INK, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Menlo"
    code.font.size = Pt(8.5)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0


def ensure_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.insert(0, abstract)
    return abstract_id


def new_numbering_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first

    if not first:
        header = section.header
        header.is_linked_to_previous = False
        for child in list(header._element):
            header._element.remove(child)
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("《秦灭六国》微信小游戏上架材料包")
        set_run_font(run, size=9, color=MUTED)
        footer = section.footer
        footer.is_linked_to_previous = False
        for child in list(footer._element):
            footer._element.remove(child)
        add_page_field(footer.add_paragraph())


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("微信小游戏上架材料包")
    set_run_font(r, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(title, "《秦灭六国》")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(subtitle, "个人开发者 · V1.0 · 提审与备案底稿")

    for _ in range(5):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    add_inline_markdown(meta, "开发者：师朋飞", MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    add_inline_markdown(meta, "生成日期：2026 年 8 月 5 日", MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_markdown(meta, "首版边界：无登录 · 无广告 · 无内购 · 无支付", MUTED)


def add_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    if any(len(row) != col_count for row in rows):
        return
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = compute_widths(rows)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline_markdown(p, text)
            for run in p.runs:
                set_run_font(run, size=9.5, bold=(r_idx == 0))
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(p, text, INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def render_markdown(doc, text, decimal_abstract_id, skip_first_h1=False):
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    first_h1_seen = False
    previous_was_numbered = False
    current_num_id = None

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            previous_was_numbered = False
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                add_inline_markdown(p, "\n".join(code_lines))
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F4F7")
                p_pr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|$", lines[i + 1].strip()):
            previous_was_numbered = False
            table_rows = []
            table_rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            previous_was_numbered = False
            level = len(heading.group(1))
            text_value = heading.group(2)
            if level == 1 and not first_h1_seen:
                first_h1_seen = True
                if skip_first_h1:
                    i += 1
                    continue
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(p, text_value)
            i += 1
            continue

        if re.match(r"^-{3,}$", line.strip()):
            previous_was_numbered = False
            i += 1
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            previous_was_numbered = False
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            if not previous_was_numbered or current_num_id is None:
                current_num_id = new_numbering_instance(doc, decimal_abstract_id)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, current_num_id)
            add_inline_markdown(p, numbered.group(1))
            previous_was_numbered = True
            i += 1
            continue

        if line.startswith(">"):
            previous_was_numbered = False
            add_callout(doc, line.lstrip("> "))
            i += 1
            continue

        if not line.strip():
            previous_was_numbered = False
            i += 1
            continue

        p = doc.add_paragraph()
        previous_was_numbered = False
        add_inline_markdown(p, line)
        i += 1


def main():
    missing = [name for name in SOURCE_FILES if not (PACK_DIR / name).exists()]
    if missing:
        raise SystemExit(f"Missing sources: {', '.join(missing)}")

    doc = Document()
    configure_styles(doc)
    decimal_abstract_id = ensure_decimal_numbering(doc)
    configure_section(doc.sections[0], first=True)
    add_cover(doc)

    for index, name in enumerate(SOURCE_FILES):
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, first=False)
        source = (PACK_DIR / name).read_text(encoding="utf-8")
        render_markdown(doc, source, decimal_abstract_id, skip_first_h1=False)

    doc.core_properties.title = "《秦灭六国》微信小游戏上架材料包 V1.0"
    doc.core_properties.subject = "个人开发者微信小游戏备案、资质、隐私与提审材料"
    doc.core_properties.author = "师朋飞"
    doc.core_properties.keywords = "微信小游戏, 上架, 备案, 隐私, 自审, 秦灭六国"
    doc.core_properties.comments = "由仓库 Markdown 底稿生成；提交时以微信后台最新模板为准。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"build failed: {exc}", file=sys.stderr)
        raise
