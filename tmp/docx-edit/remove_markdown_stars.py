from pathlib import Path
from docx import Document


SOURCE = Path(
    "/Users/zhuangxiji/Desktop/秦灭六国软著申请/"
    "Codex-最终材料-20260812/02-可编辑源文件/"
    "02-秦灭六国游戏软件-V1.0-软件设计与使用说明书.docx"
)
OUTPUT = Path("tmp/docx-edit/02-秦灭六国游戏软件-V1.0-软件设计与使用说明书.docx")
OLD = "**但 V1.0 仅作为浏览器应用发布，不作为小游戏版本发布**。"
NEW = "但 V1.0 仅作为浏览器应用发布，不作为小游戏版本发布。"


def replace_in_paragraph(paragraph) -> int:
    count = 0
    for run in paragraph.runs:
        if OLD in run.text:
            run.text = run.text.replace(OLD, NEW)
            count += 1
    return count


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


document = Document(SOURCE)
replacement_count = sum(replace_in_paragraph(p) for p in iter_paragraphs(document))
if replacement_count != 1:
    raise RuntimeError(f"Expected exactly one replacement, found {replacement_count}")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(f"Saved {OUTPUT} with {replacement_count} replacement")
